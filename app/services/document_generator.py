"""
公文Word文档生成器 - 严格按照GB/T9704-2012标准
"""
import re
import logging
from io import BytesIO
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tempfile
import os

from app.utils.text_processor import text_processor

logger = logging.getLogger(__name__)

class OfficialDocumentGenerator:
    """党政机关公文生成器"""
    
    def __init__(self):
        """初始化文档生成器"""
        self.document = Document()
        self._setup_chinese_number_mapping()
        self._setup_styles()  # 初始化样式

    def _setup_chinese_number_mapping(self):
        """设置中文数字映射"""
        self.chinese_numbers = [
            "一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
            "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十"
        ]
        
        self.chinese_sub_numbers = [
            "（一）", "（二）", "（三）", "（四）", "（五）", "（六）", "（七）", "（八）", "（九）", "（十）",
            "（十一）", "（十二）", "（十三）", "（十四）", "（十五）", "（十六）", "（十七）", "（十八）", "（十九）", "（二十）"
        ]
    
    def generate_document(self, title: str, issuing_department: str, issue_date: str, 
                         content: str, receiving_department: Optional[str] = None,
                         has_attachments: bool = False, attachments: Optional[List[Dict]] = None) -> bytes:
        """
        生成符合GB/T9704-2012标准的公文
        
        Args:
            title: 文档标题
            issuing_department: 发文部门
            issue_date: 发文日期
            content: 正文内容
            receiving_department: 收文部门
            has_attachments: 是否有附件
            attachments: 附件列表
            
        Returns:
            bytes: 生成的Word文档字节流
        """
        try:
            # 创建新文档
            self.document = Document()
            
            # 设置页面格式
            self._setup_page_format()
            
            # 设置样式
            self._setup_styles()
            # 在生成 Word 文档前，应该对内容做一次"清理和去重"，避免发文部门、发文日期、标题等信息重复出现在正文和落款。
            # 清理和处理内容，确保去除重复的标题
            cleaned_content = self._clean_content_remove_title(content, title, issuing_department, issue_date, receiving_department)
            logger.info(f"cleaned_content 长度：{len(cleaned_content)} 字符")
            # 添加文档标题
            self._add_document_title(title)
            
            # 添加空行
            self._add_empty_line()
            
            # 添加正文内容
            self._add_document_content(cleaned_content)
            
            # 添加附件说明（在正文下方，落款之前）
            if has_attachments and attachments:
                self._add_attachment_references(self.document, attachments)
            
            # 添加落款
            self._add_signature(issuing_department, issue_date)
            
            # 添加附件内容（另起页）
            if has_attachments and attachments:
                for i, attachment in enumerate(attachments, 1):
                    self._add_attachment_content(attachment, i)
            
            # 设置页码
            self._add_page_numbers()
            
            # 保存为字节流
            document_stream = BytesIO()
            self.document.save(document_stream)
            document_stream.seek(0)
            
            logger.info("公文文档生成成功")
            return document_stream.getvalue()
            
        except Exception as e:
            logger.error(f"生成文档时发生错误: {str(e)}")
            raise
    
    def _clean_content_remove_title(self, content: str, title: str, issuing_department: str, issue_date: str, receiving_department: str) -> str:
        """
        清理内容，确保去除重复的标题、发文部门、发文日期、收文部门、附件标题
        
        Args:
            content: 原始内容
            title: 文档标题
            issuing_department: 发文部门
            issue_date: 发文日期
            receiving_department: 收文部门

        Returns:
            str: 清理后的内容
        """
        try:
            lines = content.split('\n')
            cleaned_lines = []
            
            for line in lines:
                line_stripped = line.strip()
                
                # 跳过与文档标题相同的行（各种格式）
                if (line_stripped == title or 
                    line_stripped == f"# {title}" or
                    line_stripped == f"## {title}" or
                    line_stripped == f"### {title}" or
                    line_stripped == issuing_department or
                    line_stripped == issue_date or
                    line_stripped == receiving_department):
                    continue
                
                # 跳过明显重复的编号行
                if self._is_duplicate_numbering(line_stripped):
                    continue
                    
                cleaned_lines.append(line)
            
            return '\n'.join(cleaned_lines).strip()
            
        except Exception as e:
            logger.error(f"清理内容时发生错误: {str(e)}")
            return content
    
    
    def _is_duplicate_numbering(self, line: str) -> bool:
        """
        检查是否是重复的编号行
        
        Args:
            line: 要检查的行
            
        Returns:
            bool: 是否是重复编号
        """
        try:
            line = line.strip()
            
            # 检查是否有重复的编号模式，如："(一)一、"
            patterns = [
                r'（一）\s*一、',  # （一）一、
                r'（二）\s*二、',  # （二）二、
                r'（三）\s*三、',  # （三）三、
                r'（四）\s*四、',  # （四）四、
                r'（五）\s*五、',  # （五）五、
            ]
            
            import re
            for pattern in patterns:
                if re.search(pattern, line):
                    return True
            
            # 检查其他重复模式
            if re.search(r'（[一二三四五六七八九十]+）\s*[一二三四五六七八九十]+、', line):
                return True
                
            return False
            
        except Exception as e:
            logger.error(f"检查重复编号时发生错误: {str(e)}")
            return False
    
    def _setup_page_format(self):
        """设置页面格式"""
        try:
            # 获取节对象
            section = self.document.sections[0]
            
            # 设置页边距 (GB/T9704-2012: 上下2.54cm，左右3.18cm)
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.18)
            section.right_margin = Cm(3.18)
            
            # 设置页面大小为A4
            section.page_height = Cm(29.7)
            section.page_width = Cm(21)
            
            logger.info("页面格式设置完成")
            
        except Exception as e:
            logger.error(f"设置页面格式时发生错误: {str(e)}")
            raise
    
    def _setup_styles(self):
        """设置文档样式"""
        try:
            styles = self.document.styles
            
            # 1. 标题样式 (方正小标宋简体，二号，行距35磅)
            if 'DocumentTitle' not in [s.name for s in styles]:
                title_style = styles.add_style('DocumentTitle', WD_STYLE_TYPE.PARAGRAPH)
                title_font = title_style.font
                title_font.name = '方正小标宋简体'
                title_font.size = Pt(22)  # 二号字体
                title_font.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
                
                title_paragraph = title_style.paragraph_format
                title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                title_paragraph.line_spacing = Pt(35)  # 行距35磅
                title_paragraph.space_before = Pt(0)
                title_paragraph.space_after = Pt(0)
            
            # 2. 一级标题样式 (黑体，三号)
            if 'Heading1Official' not in [s.name for s in styles]:
                h1_style = styles.add_style('Heading1Official', WD_STYLE_TYPE.PARAGRAPH)
                h1_font = h1_style.font
                h1_font.name = '黑体'
                h1_font.size = Pt(16)  # 三号字体
                h1_font.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                
                h1_paragraph = h1_style.paragraph_format
                h1_paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                h1_paragraph.line_spacing = Pt(30)  # 行距30磅
                h1_paragraph.space_before = Pt(0)
                h1_paragraph.space_after = Pt(0)
                h1_paragraph.first_line_indent = Pt(32)  # 首行缩进2字符
            
            # 3. 二级标题样式 (楷体_GB2312，三号)
            if 'Heading2Official' not in [s.name for s in styles]:
                h2_style = styles.add_style('Heading2Official', WD_STYLE_TYPE.PARAGRAPH)
                h2_font = h2_style.font
                h2_font.name = '楷体_GB2312'
                h2_font.size = Pt(16)  # 三号字体
                h2_font.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
                
                h2_paragraph = h2_style.paragraph_format
                h2_paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                h2_paragraph.line_spacing = Pt(30)  # 行距30磅
                h2_paragraph.space_before = Pt(0)
                h2_paragraph.space_after = Pt(0)
                h2_paragraph.first_line_indent = Pt(32)  # 首行缩进2字符
            
            # 4. 三级标题样式 (仿宋_GB2312，三号，加粗)
            if 'Heading3Official' not in [s.name for s in styles]:
                h3_style = styles.add_style('Heading3Official', WD_STYLE_TYPE.PARAGRAPH)
                h3_font = h3_style.font
                h3_font.name = '仿宋_GB2312'
                h3_font.size = Pt(16)  # 三号字体
                h3_font.bold = True  # 加粗
                h3_font.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                
                h3_paragraph = h3_style.paragraph_format
                h3_paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                h3_paragraph.line_spacing = Pt(30)  # 行距30磅
                h3_paragraph.space_before = Pt(0)
                h3_paragraph.space_after = Pt(0)
                h3_paragraph.first_line_indent = Pt(32)  # 首行缩进2字符
            
            # 5. 正文样式 (仿宋_GB2312，三号，行距30磅)
            if 'BodyOfficial' not in [s.name for s in styles]:
                body_style = styles.add_style('BodyOfficial', WD_STYLE_TYPE.PARAGRAPH)
                body_font = body_style.font
                body_font.name = '仿宋_GB2312'
                body_font.size = Pt(16)  # 三号字体
                body_font.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                
                body_paragraph = body_style.paragraph_format
                body_paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                body_paragraph.line_spacing = Pt(30)  # 行距30磅
                body_paragraph.space_before = Pt(0)
                body_paragraph.space_after = Pt(0)
                body_paragraph.first_line_indent = Pt(32)  # 首行缩进2字符
            
            # 6. 附件样式 (黑体，三号)
            if 'AttachmentOfficial' not in [s.name for s in styles]:
                attach_style = styles.add_style('AttachmentOfficial', WD_STYLE_TYPE.PARAGRAPH)
                attach_font = attach_style.font
                attach_font.name = '黑体'
                attach_font.size = Pt(16)  # 三号字体
                attach_font.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                
                attach_paragraph = attach_style.paragraph_format
                attach_paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                attach_paragraph.line_spacing = Pt(30)  # 行距30磅
                attach_paragraph.space_before = Pt(0)
                attach_paragraph.space_after = Pt(0)
            
            # 7. 附件标题样式 (黑体，三号，居中)
            if 'AttachmentTitle' not in [s.name for s in styles]:
                att_title_style = styles.add_style('AttachmentTitle', WD_STYLE_TYPE.PARAGRAPH)
                att_title_font = att_title_style.font
                att_title_font.name = '黑体'
                att_title_font.size = Pt(16)  # 三号字体
                att_title_font.bold = True
                att_title_font.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                
                att_title_paragraph = att_title_style.paragraph_format
                att_title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                att_title_paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                att_title_paragraph.line_spacing = Pt(30)  # 行距30磅
                att_title_paragraph.space_before = Pt(0)
                att_title_paragraph.space_after = Pt(0)
                att_title_paragraph.first_line_indent = Pt(0)  # 顶格
            
            # 8. 附件内容标题样式 (黑体，三号，居中)
            if 'AttachmentContentTitle' not in [s.name for s in styles]:
                att_content_title_style = styles.add_style('AttachmentContentTitle', WD_STYLE_TYPE.PARAGRAPH)
                att_content_title_font = att_content_title_style.font
                att_content_title_font.name = '黑体'
                att_content_title_font.size = Pt(16)  # 三号字体
                att_content_title_font.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                
                att_content_title_paragraph = att_content_title_style.paragraph_format
                att_content_title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                att_content_title_paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                att_content_title_paragraph.line_spacing = Pt(30)  # 行距30磅
                att_content_title_paragraph.space_before = Pt(0)
                att_content_title_paragraph.space_after = Pt(0)
            
            logger.info("文档样式设置完成")
            
        except Exception as e:
            logger.error(f"设置文档样式时发生错误: {str(e)}")
            raise
    
    def _add_document_title(self, title: str):
        """添加文档标题"""
        try:
            # 添加标题段落
            title_paragraph = self.document.add_paragraph(style='DocumentTitle')
            
            # 添加标题文本运行对象，确保字体正确应用
            title_run = title_paragraph.add_run(title)
            
            # 强制设置字体属性
            title_run.font.name = '方正小标宋简体'
            title_run.font.size = Pt(22)  # 二号字体
            title_run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
            title_run.font.bold = False
            
            # 如果标题过长，进行换行处理，形成梯形或菱形排列
            if len(title) > 20:
                # 清除原内容
                title_paragraph.clear()
                
                # 简单的换行处理
                words = list(title)
                mid_point = len(words) // 2
                
                # 寻找合适的断点
                break_point = mid_point
                for i in range(mid_point - 3, mid_point + 4):
                    if i < len(words) and words[i] in ['的', '关于', '和', '与', '及']:
                        break_point = i + 1
                        break
                
                if break_point < len(words):
                    line1 = ''.join(words[:break_point])
                    line2 = ''.join(words[break_point:])
                    
                    # 添加第一行
                    line1_run = title_paragraph.add_run(line1)
                    line1_run.font.name = '方正小标宋简体'
                    line1_run.font.size = Pt(22)
                    line1_run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
                    
                    # 添加换行
                    title_paragraph.add_run('\n')
                    
                    # 添加第二行
                    line2_run = title_paragraph.add_run(line2)
                    line2_run.font.name = '方正小标宋简体'
                    line2_run.font.size = Pt(22)
                    line2_run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
            
            logger.info("文档标题添加完成")
            
        except Exception as e:
            logger.error(f"添加文档标题时发生错误: {str(e)}")
            raise
    
    def _add_empty_line(self):
        """添加空行"""
        try:
            empty_paragraph = self.document.add_paragraph()
            empty_paragraph.style = 'BodyOfficial'
            
        except Exception as e:
            logger.error(f"添加空行时发生错误: {str(e)}")
    
    def _add_document_content(self, content: str):
        """添加文档正文内容"""
        try:
            # 格式化内容
            formatted_content = text_processor.format_content_for_document(content)
            logger.info(f"formatted_content 长度：{len(formatted_content)} 字符")
            level1_counter = 0
            level2_counter = 0
            level3_counter = 0
            
            for item in formatted_content:
                logger.info(f"item：{item}")
                if item['type'] == 'header1':
                    # 一级标题：一、二、三、
                    level1_counter += 1
                    level2_counter = 0  # 重置二级计数
                    level3_counter = 0  # 重置三级计数
                    
                    if level1_counter <= len(self.chinese_numbers):
                        title_text = f"{self.chinese_numbers[level1_counter-1]}、{item['text']}"
                    else:
                        title_text = f"{level1_counter}、{item['text']}"
                    logger.info(f"一级标题：{title_text}")
                    paragraph = self.document.add_paragraph(style='Heading1Official')
                    run = paragraph.add_run(title_text)
                    # 显式设置字体属性
                    run.font.name = '黑体'
                    run.font.size = Pt(16)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    
                    # paragraph = self.document.add_paragraph(title_text, style='Heading1Official')
                    
                elif item['type'] == 'header2':
                    # 二级标题：（一）、（二）、（三）
                    level2_counter += 1
                    level3_counter = 0  # 重置三级计数
                    
                    if level2_counter <= len(self.chinese_sub_numbers):
                        title_text = f"{self.chinese_sub_numbers[level2_counter-1]}{item['text']}"
                    else:
                        title_text = f"（{level2_counter}）{item['text']}"
                    logger.info(f"二级标题：{title_text}")
                    paragraph = self.document.add_paragraph(style='Heading2Official')
                    run = paragraph.add_run(title_text)
                    # 显式设置字体属性
                    run.font.name = '楷体_GB2312'
                    run.font.size = Pt(16)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
                    
                elif item['type'] == 'header3':
                    # 三级标题：1.、2.、3.
                    level3_counter += 1
                    title_text = f"{level3_counter}.{item['text']}"
                    
                    paragraph = self.document.add_paragraph(style='Heading3Official')
                    run = paragraph.add_run(title_text)
                    # 显式设置字体属性
                    run.font.name = '仿宋_GB2312'   
                    run.font.size = Pt(16)
                    run.font.bold = True
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    
                else:
                    # 普通段落
                    paragraph = self.document.add_paragraph(style='BodyOfficial')
                    run = paragraph.add_run(item['text'])
                    # 显式设置字体属性
                    run.font.name = '仿宋_GB2312'
                    run.font.size = Pt(16)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            
            logger.info("文档内容添加完成")
            
        except Exception as e:
            logger.error(f"添加文档内容时发生错误: {str(e)}")
            raise
    
    def _add_attachment_references(self, doc: Document, attachments: List[Dict]) -> None:
        """
        在文档中添加附件引用，严格按照公文格式要求
        
        格式要求：
        - 正文下空一行左空二字编排"附件"二字，后标全角冒号和附件名称
        - 多个附件使用阿拉伯数字标注顺序
        - 附件名称后不加标点
        - 附件名称较长需回行时，应与上一行附件名称的首字对齐
        
        Args:
            doc: Word文档对象
            attachments: 附件列表
        """
        if not attachments:
            return
        
        # 减少日志输出，只打印关键信息
        logger.info(f"添加附件引用，共{len(attachments)}个附件")
        
        # 添加空行
        doc.add_paragraph()
        
        # 按照新的格式要求：附件：附件1、标题 附件2、标题
        if len(attachments) == 1:
            # 只有1个附件：附件：附件1、标题
            attachment_paragraph = doc.add_paragraph()
            attachment_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 左空二字（32磅）
            attachment_paragraph.paragraph_format.first_line_indent = Pt(32)
            
            # 添加"附件：附件1、"
            attachment_run = attachment_paragraph.add_run("附件：附件1、")
            attachment_run.font.name = "仿宋_GB2312"
            attachment_run.font.size = Pt(16)
            attachment_run.bold = True
            # 强制设置字体
            if attachment_run._element.rPr is not None:
                attachment_run._element.rPr.rFonts.set(qn('w:ascii'), "仿宋_GB2312")
                attachment_run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋_GB2312")
                attachment_run._element.rPr.rFonts.set(qn('w:hAnsi'), "仿宋_GB2312")
            
            # 获取附件标题
            attachment_title = attachments[0].get('extracted_title', attachments[0].get('title', attachments[0].get('name', '附件')))
            
            # 添加附件标题（不加标点）
            title_run = attachment_paragraph.add_run(attachment_title)
            title_run.font.name = "仿宋_GB2312"
            title_run.font.size = Pt(16)
            # 强制设置字体
            if title_run._element.rPr is not None:
                title_run._element.rPr.rFonts.set(qn('w:ascii'), "仿宋_GB2312")
                title_run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋_GB2312")
                title_run._element.rPr.rFonts.set(qn('w:hAnsi'), "仿宋_GB2312")
            
        else:
            # 多个附件：第一行"附件："，后续每行一个附件
            # 第一行：附件：
            first_paragraph = doc.add_paragraph()
            first_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            first_paragraph.paragraph_format.first_line_indent = Pt(32)
            
            first_run = first_paragraph.add_run("附件：")
            first_run.font.name = "仿宋_GB2312"
            first_run.font.size = Pt(16)
            first_run.bold = True
            # 强制设置字体
            if first_run._element.rPr is not None:
                first_run._element.rPr.rFonts.set(qn('w:ascii'), "仿宋_GB2312")
                first_run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋_GB2312")
                first_run._element.rPr.rFonts.set(qn('w:hAnsi'), "仿宋_GB2312")
            
            # 每个附件占一行：附件1、标题
            for i, attachment in enumerate(attachments, 1):
                attachment_paragraph = doc.add_paragraph()
                attachment_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # 与"附件："对齐，左空二字 + "附件"两字 = 左空四字（64磅）
                attachment_paragraph.paragraph_format.first_line_indent = Pt(64)
                
                # 添加"附件X、"
                number_run = attachment_paragraph.add_run(f"附件{i}、")
                number_run.font.name = "仿宋_GB2312"
                number_run.font.size = Pt(16)
                number_run.bold = True
                # 强制设置字体
                if number_run._element.rPr is not None:
                    number_run._element.rPr.rFonts.set(qn('w:ascii'), "仿宋_GB2312")
                    number_run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋_GB2312")
                    number_run._element.rPr.rFonts.set(qn('w:hAnsi'), "仿宋_GB2312")
                
                # 获取附件标题
                attachment_title = attachment.get('extracted_title', attachment.get('title', attachment.get('name', f'附件{i}')))
                
                # 添加附件标题（不加标点）
                title_run = attachment_paragraph.add_run(attachment_title)
                title_run.font.name = "仿宋_GB2312"
                title_run.font.size = Pt(16)
                # 强制设置字体
                if title_run._element.rPr is not None:
                    title_run._element.rPr.rFonts.set(qn('w:ascii'), "仿宋_GB2312")
                    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋_GB2312")
                    title_run._element.rPr.rFonts.set(qn('w:hAnsi'), "仿宋_GB2312")
    
    def _add_signature(self, issuing_department: str, issue_date: str):
        """添加落款"""
        try:
            # 添加两个空行
            self._add_empty_line()
            self._add_empty_line()
            
            # 添加发文部门（右对齐）
            dept_paragraph = self.document.add_paragraph(issuing_department, style='BodyOfficial')
            dept_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            dept_paragraph.paragraph_format.first_line_indent = Pt(0)
            
            # 添加发文日期（右对齐）
            date_paragraph = self.document.add_paragraph(issue_date, style='BodyOfficial')
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_paragraph.paragraph_format.first_line_indent = Pt(0)
            
            logger.info("落款添加完成")
            
        except Exception as e:
            logger.error(f"添加落款时发生错误: {str(e)}")
            raise
    
    def _add_attachment_pages(self, attachments: List[Dict]):
        """
        添加附件页面（另起页）
        
        Args:
            attachments: 附件列表
        """
        try:
            for i, attachment in enumerate(attachments, 1):
                # 添加分页符（另起页）
                self.document.add_page_break()
                
                # 添加"附件"标识（顶格，版心左上角第一行）
                attach_label = self.document.add_paragraph(f"附件{i}", style='AttachmentTitle')
                
                # 添加空行
                self._add_empty_line()
                
                # 添加附件标题（居中，版心第三行）
                attach_title = self.document.add_paragraph(attachment.get('name', f'附件{i}'), style='AttachmentContentTitle')
                
                # 添加空行
                self._add_empty_line()
                
                # 添加附件内容
                self._add_attachment_content(attachment, i)
            
            logger.info("附件页面添加完成")
            
        except Exception as e:
            logger.error(f"添加附件页面时发生错误: {str(e)}")
    
    def _add_attachment_content(self, attachment: Dict, attachment_number: int = 1):
        """
        添加附件内容，支持原样格式输出
        按照公文格式要求：
        - "附件"在另起页用三号黑体字顶格编排在版心左上角第一行
        - 附件标题居中编排在版心第三行
        - 附件顺序号和附件标题应当与附件说明的表述一致
        
        Args:
            attachment: 附件信息
            attachment_number: 附件序号
        """
        try:
            # 调试信息：打印附件数据结构
            logger.info(f"添加附件内容，附件类型: {attachment.get('type', 'unknown')}")
            
            attach_type = attachment.get('type', 'text')
            
            logger.info(f"开始添加附件内容，类型: {attach_type}")
            
            # 添加分页符（另起页）
            self._add_page_break()
            
            # 添加"附件"标识（顶格，版心左上角第一行，三号黑体字）
            self._add_attachment_header(attachment_number)
            
            # 添加空行（版心第二行）
            self._add_empty_line()
            
            if attach_type in ['word', 'docx']:
                # Word文档内容，如果有原始内容则直接拼接，保持原有格式
                if 'content' in attachment:
                    # 直接拼接Word文档内容，不添加额外标题
                    # Word文档中的原始标题会自动保持原有格式
                    self._merge_word_content(attachment['content'])
                else:
                    # 如果没有原始Word内容，则添加标题
                    attachment_title = attachment.get('extracted_title', attachment.get('title', attachment.get('name', '附件')))
                    self._add_attachment_title(attachment_title)
                    
                    content = attachment.get('markdown_content', '')
                    self._add_word_content(content)
            else:
                # 非Word文档，需要添加标题
                attachment_title = attachment.get('extracted_title', attachment.get('title', attachment.get('name', '附件')))
                self._add_attachment_title(attachment_title)
                
                if attach_type == 'csv' or attach_type == 'excel':
                    content = attachment.get('markdown_content', '')
                    self._add_table_from_markdown(content)
                else:
                    # 普通文本内容
                    content = attachment.get('markdown_content', '')
                    self._add_text_content(content)
            
        except Exception as e:
            logger.error(f"添加附件内容时发生错误: {str(e)}")
    
    def _add_attachment_header(self, attachment_number: int = 1):
        """
        添加"附件"标识，按照公文格式要求：
        - 用三号黑体字顶格编排在版心左上角第一行
        - 多个附件时标记序号
        
        Args:
            attachment_number: 附件序号
        """
        try:
            # 添加"附件"标识段落
            header_paragraph = self.document.add_paragraph()
            header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 顶格编排（无缩进）
            header_paragraph.paragraph_format.first_line_indent = Pt(0)
            header_paragraph.paragraph_format.left_indent = Pt(0)
            
            # 添加"附件"文字和序号
            header_text = f"附件{attachment_number}"
            header_run = header_paragraph.add_run(header_text)
            
            # 设置字体属性：三号黑体字
            header_run.font.name = '黑体'
            header_run.font.size = Pt(16)  # 三号字体
            header_run.font.bold = True
            # 强制设置字体
            if header_run._element.rPr is not None:
                header_run._element.rPr.rFonts.set(qn('w:ascii'), '黑体')
                header_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                header_run._element.rPr.rFonts.set(qn('w:hAnsi'), '黑体')
            
        except Exception as e:
            logger.error(f"添加附件标识时发生错误: {str(e)}")
    
    def _add_attachment_title(self, title: str):
        """
        添加附件标题，按照公文格式要求：
        - 居中编排在版心第三行
        
        Args:
            title: 附件标题
        """
        try:
            # 添加附件标题段落
            title_paragraph = self.document.add_paragraph()
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加标题文字
            title_run = title_paragraph.add_run(title)
            
            # 设置字体属性：三号黑体字
            title_run.font.name = '黑体'
            title_run.font.size = Pt(16)  # 三号字体
            title_run.font.bold = True
            # 强制设置字体
            if title_run._element.rPr is not None:
                title_run._element.rPr.rFonts.set(qn('w:ascii'), '黑体')
                title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                title_run._element.rPr.rFonts.set(qn('w:hAnsi'), '黑体')
            
            # 添加空行
            self._add_empty_line()
            
        except Exception as e:
            logger.error(f"添加附件标题时发生错误: {str(e)}")
    
    def _add_page_break(self):
        """
        添加分页符
        """
        try:
            # 添加分页符
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)
            
        except Exception as e:
            logger.error(f"添加分页符时发生错误: {str(e)}")
    
    def _merge_word_content(self, word_content: bytes):
        """
        合并Word文档内容到当前文档
        增强版本，支持直接docx拼接的备选方案
        """
        try:
            # 方法1：尝试精确复制格式
            attachment_doc = Document(BytesIO(word_content))
            
            # 检查是否包含复杂表格
            has_complex_tables = self._has_complex_tables(attachment_doc)
            logger.info(f"复杂表格检测结果: {has_complex_tables}, 表格数量: {len(attachment_doc.tables)}")
            
            if has_complex_tables:
                logger.info("检测到复杂表格，使用docx直接拼接方案")
                success = self._try_direct_docx_merge(word_content)
                if success:
                    return
                else:
                    logger.warning("docx直接拼接失败，回退到元素复制方案")
            
            # 方法2：逐个处理段落和表格，保持原有顺序
            self._extract_and_add_content(attachment_doc)
            
            logger.info("Word文档拼接完成")
            
        except Exception as e:
            logger.error(f"拼接Word文档时发生错误: {str(e)}")
            # 如果拼接失败，尝试提取文本内容
            try:
                attachment_doc = Document(BytesIO(word_content))
                self._extract_and_add_content(attachment_doc)
            except Exception as e2:
                logger.error(f"提取Word文档文本内容也失败: {str(e2)}")
    
    def _has_complex_tables(self, doc):
        """
        检查文档是否包含复杂表格
        """
        try:
            for table in doc.tables:
                # 检查是否有合并单元格
                for row in table.rows:
                    for cell in row.cells:
                        try:
                            tc = cell._tc
                            tcPr = tc.tcPr
                            if tcPr is not None:
                                # 检查水平合并
                                gridSpan = tcPr.find(qn('w:gridSpan'))
                                if gridSpan is not None:
                                    span = int(gridSpan.get(qn('w:val'), 1))
                                    if span > 1:
                                        return True
                                
                                # 检查垂直合并
                                vMerge = tcPr.find(qn('w:vMerge'))
                                if vMerge is not None:
                                    return True
                        except:
                            continue
                
                # 检查行数是否不一致（可能存在合并单元格）
                if len(table.rows) > 1:
                    first_row_cells = len(table.rows[0].cells)
                    for row in table.rows[1:]:
                        if len(row.cells) != first_row_cells:
                            return True
            
            return False
        except:
            return False
    
    def _try_direct_docx_merge(self, word_content: bytes):
        """
        尝试直接拼接docx文档
        使用XML级别的直接合并来确保原样输出，同时修复附件位置问题
        """
        try:
            import tempfile
            import os
            import copy
            
            # 创建临时文件
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(word_content)
                temp_path = temp_file.name
            
            try:
                # 读取附件文档
                attachment_doc = Document(temp_path)
                
                # 调试信息：打印附件文档的基本信息
                logger.info(f"附件文档信息: 段落数={len(attachment_doc.paragraphs)}, 表格数={len(attachment_doc.tables)}")
                
                # 获取当前文档的body
                current_body = self.document.element.body
                
                # 获取附件文档的body
                attachment_body = attachment_doc.element.body
                
                # 记录当前文档的段落数（用于调试）
                original_para_count = len(current_body.xpath('.//w:p'))
                
                # 立即复制附件文档的所有元素到当前位置
                # 这样可以确保内容出现在正确的附件标题下方
                copied_count = 0
                for element in attachment_body:
                    if element.tag.endswith('sectPr'):
                        # 跳过节属性，避免影响文档结构
                        continue
                    
                    # 深度复制元素并立即添加到当前文档
                    new_element = copy.deepcopy(element)
                    current_body.append(new_element)
                    copied_count += 1
                
                # 记录复制后的段落数（用于调试）
                final_para_count = len(current_body.xpath('.//w:p'))
                
                logger.info(f"docx直接合并成功: 原始段落数={original_para_count}, 复制元素数={copied_count}, 最终段落数={final_para_count}")
                return True
                
            finally:
                # 清理临时文件
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
                    
        except Exception as e:
            logger.error(f"docx直接合并失败: {str(e)}")
            return False
    def _add_table_with_format(self, source_table):
        """
        添加表格并尝试保持格式
        
        Args:
            source_table: 源表格对象
        """
        try:
            # 获取表格数据和格式
            rows_data = []
            cell_formats = []
            
            for row_idx, row in enumerate(source_table.rows):
                row_data = []
                row_formats = []
                
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                    
                    # 获取单元格格式信息
                    cell_format = {
                        'paragraphs': [],
                        'width': None,
                        'vertical_alignment': None,
                        'is_merged': False,
                        'grid_span': 1,
                        'v_merge': None
                    }
                    
                    # 收集段落信息
                    for para in cell.paragraphs:
                        para_info = {
                            'text': para.text,
                            'alignment': para.alignment,
                            'runs': []
                        }
                        
                        # 收集run信息
                        for run in para.runs:
                            if run.text:
                                run_info = {
                                    'text': run.text,
                                    'font_name': run.font.name,
                                    'font_size': run.font.size,
                                    'bold': run.font.bold,
                                    'italic': run.font.italic,
                                    'underline': run.font.underline
                                }
                                para_info['runs'].append(run_info)
                        
                        cell_format['paragraphs'].append(para_info)
                    
                    # 获取单元格宽度
                    try:
                        cell_format['width'] = cell.width
                    except:
                        pass
                    
                    # 获取垂直对齐方式
                    try:
                        cell_format['vertical_alignment'] = cell.vertical_alignment
                    except:
                        pass
                    
                    # 检查单元格合并信息
                    try:
                        # 检查水平合并（gridSpan）
                        tc = cell._tc
                        tcPr = tc.tcPr
                        if tcPr is not None:
                            gridSpan = tcPr.find(qn('w:gridSpan'))
                            if gridSpan is not None:
                                cell_format['grid_span'] = int(gridSpan.get(qn('w:val'), 1))
                                cell_format['is_merged'] = True
                            
                            # 检查垂直合并（vMerge）
                            vMerge = tcPr.find(qn('w:vMerge'))
                            if vMerge is not None:
                                cell_format['v_merge'] = vMerge.get(qn('w:val'), 'continue')
                                cell_format['is_merged'] = True
                    except Exception as merge_error:
                        logger.warning(f"检查单元格合并信息时发生错误: {str(merge_error)}")
                    
                    row_formats.append(cell_format)
                
                rows_data.append(row_data)
                cell_formats.append(row_formats)
            
            if not rows_data:
                return
            
            # 创建新表格
            table = self.document.add_table(rows=len(rows_data), cols=len(rows_data[0]))
            
            # 设置表格样式 - 使用更接近原始格式的样式
            table.style = 'Table Grid'
            
            # 尝试复制表格整体格式
            try:
                if hasattr(source_table, 'alignment'):
                    table.alignment = source_table.alignment
                
                # 复制表格宽度
                if hasattr(source_table, 'width'):
                    table.width = source_table.width
                    
                # 设置表格边框
                from docx.oxml import OxmlElement
                
                # 设置表格边框样式
                tbl = table._tbl
                tblPr = tbl.tblPr
                
                # 添加边框设置
                tblBorders = OxmlElement('w:tblBorders')
                
                # 设置所有边框
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tblBorders.append(border)
                
                tblPr.append(tblBorders)
                
            except Exception as border_error:
                logger.warning(f"设置表格边框时发生错误: {str(border_error)}")
            
            # 填充表格数据和格式
            for i, (row_data, row_formats) in enumerate(zip(rows_data, cell_formats)):
                for j, (cell_text, cell_format) in enumerate(zip(row_data, row_formats)):
                    if j < len(table.rows[i].cells):
                        cell = table.rows[i].cells[j]
                        
                        # 清空默认段落
                        cell.paragraphs[0].clear()
                        
                        # 设置单元格垂直对齐方式
                        try:
                            if cell_format['vertical_alignment'] is not None:
                                cell.vertical_alignment = cell_format['vertical_alignment']
                        except:
                            pass
                        
                        # 设置单元格宽度
                        try:
                            if cell_format['width'] is not None:
                                cell.width = cell_format['width']
                        except:
                            pass
                        
                        # 处理单元格合并
                        try:
                            if cell_format['is_merged']:
                                tc = cell._tc
                                tcPr = tc.tcPr
                                if tcPr is None:
                                    tcPr = OxmlElement('w:tcPr')
                                    tc.append(tcPr)
                                
                                # 设置水平合并
                                if cell_format['grid_span'] > 1:
                                    gridSpan = OxmlElement('w:gridSpan')
                                    gridSpan.set(qn('w:val'), str(cell_format['grid_span']))
                                    tcPr.append(gridSpan)
                                
                                # 设置垂直合并
                                if cell_format['v_merge'] is not None:
                                    vMerge = OxmlElement('w:vMerge')
                                    if cell_format['v_merge'] != 'continue':
                                        vMerge.set(qn('w:val'), cell_format['v_merge'])
                                    tcPr.append(vMerge)
                        except Exception as merge_error:
                            logger.warning(f"设置单元格合并时发生错误: {str(merge_error)}")
                        
                        # 设置单元格边框
                        try:
                            tc = cell._tc
                            tcPr = tc.tcPr
                            if tcPr is None:
                                tcPr = OxmlElement('w:tcPr')
                                tc.append(tcPr)
                            
                            # 添加单元格边框
                            tcBorders = OxmlElement('w:tcBorders')
                            for border_name in ['top', 'left', 'bottom', 'right']:
                                border = OxmlElement(f'w:{border_name}')
                                border.set(qn('w:val'), 'single')
                                border.set(qn('w:sz'), '4')
                                border.set(qn('w:space'), '0')
                                border.set(qn('w:color'), '000000')
                                tcBorders.append(border)
                            tcPr.append(tcBorders)
                        except Exception as cell_border_error:
                            logger.warning(f"设置单元格边框时发生错误: {str(cell_border_error)}")
                        
                        # 添加段落内容
                        for para_idx, para_info in enumerate(cell_format['paragraphs']):
                            if para_idx == 0:
                                # 使用第一个段落
                                para = cell.paragraphs[0]
                            else:
                                # 添加新段落
                                para = cell.add_paragraph()
                            
                            # 设置段落对齐方式
                            if para_info['alignment'] is not None:
                                para.alignment = para_info['alignment']
                            
                            # 设置段落格式（行间距等）
                            try:
                                pf = para.paragraph_format
                                # 设置行间距为单倍行距
                                pf.line_spacing = 1.0
                                # 设置段前段后间距
                                pf.space_before = Pt(0)
                                pf.space_after = Pt(0)
                            except Exception as pf_error:
                                logger.warning(f"设置段落格式时发生错误: {str(pf_error)}")
                            
                            # 添加runs
                            if para_info['runs']:
                                for run_info in para_info['runs']:
                                    run = para.add_run(run_info['text'])
                                    
                                    # 设置字体格式
                                    if run_info['font_name']:
                                        run.font.name = run_info['font_name']
                                    else:
                                        run.font.name = '仿宋_GB2312'
                                    
                                    if run_info['font_size']:
                                        run.font.size = run_info['font_size']
                                    else:
                                        run.font.size = Pt(16)
                                    
                                    if run_info['bold'] is not None:
                                        run.font.bold = run_info['bold']
                                    if run_info['italic'] is not None:
                                        run.font.italic = run_info['italic']
                                    if run_info['underline'] is not None:
                                        run.font.underline = run_info['underline']
                                    
                                    # 强制设置字体
                                    if run._element.rPr is not None:
                                        font_name = run_info['font_name'] if run_info['font_name'] else '仿宋_GB2312'
                                        run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
                                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                                        run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
                            else:
                                # 如果没有runs，直接添加文本
                                if para_info['text']:
                                    run = para.add_run(para_info['text'])
                                    run.font.name = '仿宋_GB2312'
                                    run.font.size = Pt(16)
                                    if run._element.rPr is not None:
                                        run._element.rPr.rFonts.set(qn('w:ascii'), '仿宋_GB2312')
                                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                                        run._element.rPr.rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
            
            # 添加表格后的空行
            self._add_empty_line()
            
        except Exception as e:
            logger.error(f"添加表格时发生错误: {str(e)}")
            # 如果复杂格式复制失败，回退到简单方法
            try:
                self._add_simple_table(source_table)
            except Exception as e2:
                logger.error(f"简单表格添加也失败: {str(e2)}")
    
    def _add_simple_table(self, source_table):
        """
        添加简单表格（回退方法）
        
        Args:
            source_table: 源表格对象
        """
        try:
            # 获取表格数据
            rows_data = []
            for row in source_table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                rows_data.append(row_data)
            
            if not rows_data:
                return
            
            # 创建新表格
            table = self.document.add_table(rows=len(rows_data), cols=len(rows_data[0]))
            table.style = 'Table Grid'
            
            # 填充表格数据
            for i, row_data in enumerate(rows_data):
                for j, cell_text in enumerate(row_data):
                    if j < len(table.rows[i].cells):
                        cell = table.rows[i].cells[j]
                        cell.text = cell_text
                        
                        # 设置单元格字体
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = '仿宋_GB2312'
                                run.font.size = Pt(16)
                                if run._element.rPr is not None:
                                    run._element.rPr.rFonts.set(qn('w:ascii'), '仿宋_GB2312')
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                                    run._element.rPr.rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
            
            # 添加表格后的空行
            self._add_empty_line()
            
        except Exception as e:
            logger.error(f"添加简单表格时发生错误: {str(e)}")
    
    def _extract_and_add_content(self, source_doc: Document):
        """
        提取文档内容并添加到当前文档
        
        Args:
            source_doc: 源文档对象
        """
        try:
            # 提取段落内容
            for para in source_doc.paragraphs:
                text = para.text.strip()
                if text:
                    new_para = self.document.add_paragraph(text, style='BodyOfficial')
                    for run in new_para.runs:
                        run.font.name = '仿宋_GB2312'
                        run.font.size = Pt(16)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            
            # 提取表格内容
            for table in source_doc.tables:
                self._add_table_with_format(table)
            
        except Exception as e:
            logger.error(f"提取文档内容时发生错误: {str(e)}")
    
    def _copy_paragraph_with_format(self, source_para_element):
        """
        复制段落元素到当前文档，保持格式
        
        Args:
            source_para_element: 源段落元素
        """
        try:
            # 在当前文档中创建新段落
            new_para = self.document.add_paragraph()
            
            # 复制段落的XML内容，保持完整格式
            new_para._element.clear()
            
            # 深度复制元素
            import copy
            copied_element = copy.deepcopy(source_para_element)
            
            # 将复制的元素添加到新段落
            for child in copied_element:
                new_para._element.append(child)
            
        except Exception as e:
            logger.error(f"复制段落时发生错误: {str(e)}")
            # 如果复制失败，尝试提取文本内容
            try:
                from docx.oxml import parse_xml
                text_content = source_para_element.text if hasattr(source_para_element, 'text') else ''
                if text_content:
                    fallback_para = self.document.add_paragraph(text_content, style='BodyOfficial')
            except:
                pass
    
    def _copy_table_with_full_format(self, source_table):
        """
        完整复制表格格式，包括单元格合并、边框、字体等
        增强版本，修复单元格格式错误和列间距问题
        """
        try:
            # 分析源表格的真实结构
            table_structure = self._analyze_table_structure(source_table)
            
            # 获取源表格的行数和实际最大列数
            source_rows = len(source_table.rows)
            max_cols = max(len(row.cells) for row in source_table.rows) if source_table.rows else 0
            
            # 创建新表格
            table = self.document.add_table(rows=source_rows, cols=max_cols)
            table.style = 'Table Grid'
            
            # 复制表格整体属性
            self._copy_table_properties(source_table, table)
            
            # 逐行复制数据和格式
            for i, source_row in enumerate(source_table.rows):
                target_row = table.rows[i]
                
                # 复制行高
                self._copy_row_properties(source_row, target_row)
                
                # 复制单元格，考虑实际的单元格数量
                actual_cells = len(source_row.cells)
                for j in range(max_cols):
                    if j < len(target_row.cells):
                        target_cell = target_row.cells[j]
                        
                        if j < actual_cells:
                            # 有对应的源单元格
                            source_cell = source_row.cells[j]
                            self._copy_cell_with_enhanced_format(source_cell, target_cell)
                        else:
                            # 没有对应的源单元格，清空目标单元格
                            target_cell.paragraphs[0].clear()
                            # 设置为空单元格的默认格式
                            self._set_empty_cell_format(target_cell)
            
            # 应用表格结构调整
            self._apply_table_structure_adjustments(table, table_structure)
            
            logger.info(f"表格复制完成，{source_rows}行 x {max_cols}列")
            return table
            
        except Exception as e:
            logger.error(f"复制表格时发生错误: {str(e)}")
            # 降级到简单复制
            return self._add_table_with_format(source_table)
    
    def _analyze_table_structure(self, source_table):
        """
        分析表格结构，识别合并单元格和实际布局
        """
        structure = {
            'merged_cells': [],
            'column_widths': [],
            'row_heights': [],
            'has_merged_cells': False
        }
        
        try:
            # 分析每一行的单元格数量和合并情况
            for i, row in enumerate(source_table.rows):
                row_info = {
                    'row_index': i,
                    'cell_count': len(row.cells),
                    'merged_cells': []
                }
                
                for j, cell in enumerate(row.cells):
                    # 检查单元格合并信息
                    try:
                        tc = cell._tc
                        tcPr = tc.tcPr
                        if tcPr is not None:
                            # 检查水平合并
                            gridSpan = tcPr.find(qn('w:gridSpan'))
                            if gridSpan is not None:
                                span = int(gridSpan.get(qn('w:val'), 1))
                                if span > 1:
                                    row_info['merged_cells'].append({
                                        'col_index': j,
                                        'span': span,
                                        'type': 'horizontal'
                                    })
                                    structure['has_merged_cells'] = True
                            
                            # 检查垂直合并
                            vMerge = tcPr.find(qn('w:vMerge'))
                            if vMerge is not None:
                                merge_type = vMerge.get(qn('w:val'), 'continue')
                                row_info['merged_cells'].append({
                                    'col_index': j,
                                    'merge_type': merge_type,
                                    'type': 'vertical'
                                })
                                structure['has_merged_cells'] = True
                    except Exception as e:
                        logger.warning(f"分析单元格合并信息时发生错误: {str(e)}")
                
                structure['merged_cells'].append(row_info)
            
            # 分析列宽
            if source_table.columns:
                for col in source_table.columns:
                    try:
                        structure['column_widths'].append(col.width)
                    except:
                        structure['column_widths'].append(None)
            
        except Exception as e:
            logger.warning(f"分析表格结构时发生错误: {str(e)}")
        
        return structure
    
    def _copy_table_properties(self, source_table, target_table):
        """
        复制表格整体属性
        """
        try:
            if hasattr(source_table, 'alignment'):
                target_table.alignment = source_table.alignment
            if hasattr(source_table, 'width'):
                target_table.width = source_table.width
            
            # 复制表格边框设置
            source_tbl = source_table._tbl
            target_tbl = target_table._tbl
            
            source_tblPr = source_tbl.tblPr
            target_tblPr = target_tbl.tblPr
            
            if source_tblPr is not None:
                # 复制表格边框
                source_borders = source_tblPr.find(qn('w:tblBorders'))
                if source_borders is not None:
                    # 移除目标表格的现有边框
                    existing_borders = target_tblPr.find(qn('w:tblBorders'))
                    if existing_borders is not None:
                        target_tblPr.remove(existing_borders)
                    
                    # 复制边框设置
                    import copy
                    new_borders = copy.deepcopy(source_borders)
                    target_tblPr.append(new_borders)
                
                # 复制表格宽度设置
                source_width = source_tblPr.find(qn('w:tblW'))
                if source_width is not None:
                    existing_width = target_tblPr.find(qn('w:tblW'))
                    if existing_width is not None:
                        target_tblPr.remove(existing_width)
                    
                    import copy
                    new_width = copy.deepcopy(source_width)
                    target_tblPr.append(new_width)
                
        except Exception as e:
            logger.warning(f"复制表格属性时发生错误: {str(e)}")
    
    def _copy_row_properties(self, source_row, target_row):
        """
        复制行属性
        """
        try:
            if hasattr(source_row, 'height'):
                target_row.height = source_row.height
            
            # 复制行的其他属性
            source_tr = source_row._tr
            target_tr = target_row._tr
            
            source_trPr = source_tr.trPr
            if source_trPr is not None:
                target_trPr = target_tr.trPr
                if target_trPr is None:
                    target_trPr = OxmlElement('w:trPr')
                    target_tr.insert(0, target_trPr)
                
                # 复制行高设置
                source_height = source_trPr.find(qn('w:trHeight'))
                if source_height is not None:
                    existing_height = target_trPr.find(qn('w:trHeight'))
                    if existing_height is not None:
                        target_trPr.remove(existing_height)
                    
                    import copy
                    new_height = copy.deepcopy(source_height)
                    target_trPr.append(new_height)
                
        except Exception as e:
            logger.warning(f"复制行属性时发生错误: {str(e)}")
    
    def _copy_cell_with_enhanced_format(self, source_cell, target_cell):
        """
        增强的单元格复制功能，修复格式错误
        """
        try:
            # 清空目标单元格
            target_cell.paragraphs[0].clear()
            
            # 复制单元格内容
            for para_idx, source_para in enumerate(source_cell.paragraphs):
                if para_idx == 0:
                    target_para = target_cell.paragraphs[0]
                else:
                    target_para = target_cell.add_paragraph()
                
                # 复制段落格式
                self._copy_paragraph_format(source_para, target_para)
                
                # 复制文本内容
                if source_para.runs:
                    for run in source_para.runs:
                        if run.text:
                            target_run = target_para.add_run(run.text)
                            self._copy_run_format(run, target_run)
                else:
                    # 如果没有runs但有文本，直接添加
                    text = source_para.text
                    if text:
                        target_run = target_para.add_run(text)
                        target_run.font.name = '仿宋_GB2312'
                        target_run.font.size = Pt(16)
                        self._set_font_family(target_run, '仿宋_GB2312')
            
            # 复制单元格属性（宽度、对齐等）
            self._copy_cell_properties(source_cell, target_cell)
            
        except Exception as e:
            logger.warning(f"复制单元格内容时发生错误: {str(e)}")
    
    def _copy_paragraph_format(self, source_para, target_para):
        """
        复制段落格式
        """
        try:
            if source_para.alignment is not None:
                target_para.alignment = source_para.alignment
            
            # 复制段落格式属性
            source_pf = source_para.paragraph_format
            target_pf = target_para.paragraph_format
            
            if source_pf.line_spacing is not None:
                target_pf.line_spacing = source_pf.line_spacing
            if source_pf.space_before is not None:
                target_pf.space_before = source_pf.space_before
            if source_pf.space_after is not None:
                target_pf.space_after = source_pf.space_after
            if source_pf.left_indent is not None:
                target_pf.left_indent = source_pf.left_indent
            if source_pf.right_indent is not None:
                target_pf.right_indent = source_pf.right_indent
            if source_pf.first_line_indent is not None:
                target_pf.first_line_indent = source_pf.first_line_indent
                
        except Exception as e:
            logger.warning(f"复制段落格式时发生错误: {str(e)}")
    
    def _copy_run_format(self, source_run, target_run):
        """
        复制文本格式
        """
        try:
            source_font = source_run.font
            target_font = target_run.font
            
            if source_font.name:
                target_font.name = source_font.name
            else:
                target_font.name = '仿宋_GB2312'
            
            if source_font.size:
                target_font.size = source_font.size
            else:
                target_font.size = Pt(16)
            
            if source_font.bold is not None:
                target_font.bold = source_font.bold
            if source_font.italic is not None:
                target_font.italic = source_font.italic
            if source_font.underline is not None:
                target_font.underline = source_font.underline
            if source_font.color.rgb is not None:
                target_font.color.rgb = source_font.color.rgb
            
            # 设置字体族
            font_name = source_font.name if source_font.name else '仿宋_GB2312'
            self._set_font_family(target_run, font_name)
            
        except Exception as e:
            logger.warning(f"复制文本格式时发生错误: {str(e)}")
            # 设置默认格式
            target_run.font.name = '仿宋_GB2312'
            target_run.font.size = Pt(16)
            self._set_font_family(target_run, '仿宋_GB2312')
    
    def _set_font_family(self, run, font_name):
        """
        安全地设置字体族
        """
        try:
            if run._element.rPr is not None:
                # 确保rFonts元素存在
                rFonts = run._element.rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    run._element.rPr.append(rFonts)
                
                # 设置各种字体属性
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:eastAsia'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rFonts.set(qn('w:cs'), font_name)
        except Exception as e:
            logger.warning(f"设置字体族时发生错误: {str(e)}")
    
    def _copy_cell_properties(self, source_cell, target_cell):
        """
        复制单元格属性，避免格式错误
        """
        try:
            # 复制单元格宽度
            if hasattr(source_cell, 'width') and source_cell.width is not None:
                target_cell.width = source_cell.width
            
            # 复制垂直对齐
            if hasattr(source_cell, 'vertical_alignment') and source_cell.vertical_alignment is not None:
                target_cell.vertical_alignment = source_cell.vertical_alignment
            
            # 安全地复制单元格XML属性
            source_tc = source_cell._tc
            target_tc = target_cell._tc
            
            source_tcPr = source_tc.tcPr
            if source_tcPr is not None:
                target_tcPr = target_tc.tcPr
                if target_tcPr is None:
                    target_tcPr = OxmlElement('w:tcPr')
                    target_tc.append(target_tcPr)
                
                # 只复制安全的属性
                safe_elements = ['w:tcW', 'w:gridSpan', 'w:vMerge', 'w:vAlign', 'w:tcBorders', 'w:shd']
                
                for element_name in safe_elements:
                    source_element = source_tcPr.find(qn(element_name))
                    if source_element is not None:
                        # 移除目标中的现有元素
                        existing_element = target_tcPr.find(qn(element_name))
                        if existing_element is not None:
                            target_tcPr.remove(existing_element)
                        
                        # 创建新元素并复制属性
                        new_element = OxmlElement(element_name)
                        
                        # 复制文本内容
                        if source_element.text:
                            new_element.text = source_element.text
                        
                        # 复制属性
                        for attr_name, attr_value in source_element.attrib.items():
                            # 过滤掉可能导致问题的属性
                            if not attr_name.startswith('{http'):
                                new_element.set(attr_name, attr_value)
                        
                        # 复制子元素
                        for child in source_element:
                            try:
                                import copy
                                new_child = copy.deepcopy(child)
                                new_element.append(new_child)
                            except Exception as child_error:
                                logger.warning(f"复制子元素时发生错误: {str(child_error)}")
                        
                        target_tcPr.append(new_element)
                
        except Exception as e:
            logger.warning(f"复制单元格属性时发生错误: {str(e)}")
    
    def _set_empty_cell_format(self, cell):
        """
        为空单元格设置默认格式
        """
        try:
            # 确保有一个段落
            if not cell.paragraphs:
                cell.add_paragraph()
            
            para = cell.paragraphs[0]
            para.clear()
            
            # 添加空的run以保持格式
            run = para.add_run('')
            run.font.name = '仿宋_GB2312'
            run.font.size = Pt(16)
            self._set_font_family(run, '仿宋_GB2312')
            
        except Exception as e:
            logger.warning(f"设置空单元格格式时发生错误: {str(e)}")
    
    def _apply_table_structure_adjustments(self, table, structure):
        """
        应用表格结构调整
        """
        try:
            # 如果有列宽信息，应用列宽
            if structure['column_widths'] and len(structure['column_widths']) > 0:
                for i, width in enumerate(structure['column_widths']):
                    if width is not None and i < len(table.columns):
                        try:
                            table.columns[i].width = width
                        except:
                            pass
            
            # 应用其他结构调整
            if structure['has_merged_cells']:
                logger.info("表格包含合并单元格，已尝试保持原始结构")
            
        except Exception as e:
            logger.warning(f"应用表格结构调整时发生错误: {str(e)}")
    
    def _copy_table_with_format(self, source_table_element):
        """
        复制表格元素到当前文档，保持格式
        
        Args:
            source_table_element: 源表格元素
        """
        try:
            # 创建表格容器段落
            table_para = self.document.add_paragraph()
            
            # 清空段落内容
            table_para._element.clear()
            
            # 深度复制表格元素
            import copy
            copied_table = copy.deepcopy(source_table_element)
            
            # 将复制的表格添加到段落
            table_para._element.append(copied_table)
            
        except Exception as e:
            logger.error(f"复制表格时发生错误: {str(e)}")
            # 如果复制失败，尝试提取表格内容并重新创建
            try:
                self._extract_and_recreate_table(source_table_element)
            except:
                pass
    
    def _copy_element_with_format(self, source_element):
        """
        复制其他元素到当前文档
        
        Args:
            source_element: 源元素
        """
        try:
            # 创建一个容器段落
            container_para = self.document.add_paragraph()
            
            # 清空段落内容
            container_para._element.clear()
            
            # 深度复制元素
            import copy
            copied_element = copy.deepcopy(source_element)
            
            # 将复制的元素添加到容器
            container_para._element.append(copied_element)
            
        except Exception as e:
            logger.error(f"复制元素时发生错误: {str(e)}")
    
    def _extract_and_recreate_table(self, source_table_element):
        """
        提取表格内容并重新创建表格
        
        Args:
            source_table_element: 源表格元素
        """
        try:
            # 这里可以添加表格提取和重建逻辑
            # 如果需要更复杂的表格处理，可以在这里实现
            pass
        except Exception as e:
            logger.error(f"重新创建表格时发生错误: {str(e)}")
    
    def _add_page_numbers(self):
        """
        为整个文档添加连续页码
        """
        try:
            # 获取文档的第一个节
            section = self.document.sections[0]
            
            # 获取页脚
            footer = section.footer
            
            # 清空现有页脚内容
            footer._element.clear()
            
            # 创建页脚段落
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加页码字段
            run = footer_para.add_run()
            
            # 创建页码字段的XML
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            # 添加到run中
            run._element.append(fldChar1)
            run._element.append(instrText)
            run._element.append(fldChar2)
            
            # 设置页码字体
            run.font.name = '仿宋_GB2312'
            run.font.size = Pt(14)
            
            logger.info("页码添加完成")
            
        except Exception as e:
            logger.error(f"添加页码时发生错误: {str(e)}")
    
    def _add_word_content(self, content: str):
        """
        添加Word文档内容，保持原样格式
        
        Args:
            content: Word文档提取的内容
        """
        try:
            # 按段落分割内容
            paragraphs = content.split('\n\n')
            
            for paragraph_content in paragraphs:
                paragraph_content = paragraph_content.strip()
                if not paragraph_content:
                    continue
                
                # 检查是否是表格内容
                if '|' in paragraph_content and '---' in paragraph_content:
                    # 表格内容，单独处理
                    self._add_table_from_markdown(paragraph_content)
                else:
                    # 普通段落内容
                    lines = paragraph_content.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:
                            # 检查是否是标题行（简单判断）
                            if self._is_title_line(line):
                                # 使用附件标题样式
                                paragraph = self.document.add_paragraph(style='AttachmentTitle')
                                run = paragraph.add_run(line)
                                run.font.name = '黑体'
                                run.font.size = Pt(16)
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                            else:
                                # 使用正文样式
                                paragraph = self.document.add_paragraph(style='BodyOfficial')
                                run = paragraph.add_run(line)
                                run.font.name = '仿宋_GB2312'
                                run.font.size = Pt(16)
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            
        except Exception as e:
            logger.error(f"添加Word内容时发生错误: {str(e)}")
    
    def _add_text_content(self, content: str):
        """
        添加普通文本内容
        
        Args:
            content: 文本内容
        """
        try:
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    paragraph = self.document.add_paragraph(style='BodyOfficial')
                    run = paragraph.add_run(line)
                    run.font.name = '仿宋_GB2312'
                    run.font.size = Pt(16)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            
        except Exception as e:
            logger.error(f"添加文本内容时发生错误: {str(e)}")
    
    def _is_title_line(self, line: str) -> bool:
        """
        判断是否是标题行
        
        Args:
            line: 文本行
            
        Returns:
            bool: 是否是标题行
        """
        # 简单的标题判断规则
        title_patterns = [
            r'^第[一二三四五六七八九十]+章',  # 第X章
            r'^第[一二三四五六七八九十]+节',  # 第X节
            r'^[一二三四五六七八九十]+、',    # 一、二、三、
            r'^（[一二三四五六七八九十]+）',  # （一）（二）（三）
            r'^\d+\.',                      # 1. 2. 3.
            r'^[A-Z]+\.',                   # A. B. C.
        ]
        
        for pattern in title_patterns:
            if re.match(pattern, line):
                return True
        
        # 如果行长度较短且不包含标点符号，可能是标题
        if len(line) < 20 and not any(punct in line for punct in ['。', '，', '；', '：', '？', '！']):
            return True
        
        return False
    
    def _add_table_from_markdown(self, markdown_table: str):
        """
        从markdown表格创建Word表格
        
        Args:
            markdown_table: markdown格式的表格
        """
        try:
            lines = [line.strip() for line in markdown_table.split('\n') if line.strip()]
            if len(lines) < 2:
                return
            
            # 解析表格头部
            headers = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
            if not headers:
                return
            
            # 跳过分隔线
            data_lines = [line for line in lines[2:] if '|' in line]
            
            # 创建表格
            table = self.document.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            # 设置表头
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                if i < len(hdr_cells):
                    hdr_cells[i].text = header
                    # 设置表头字体
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '仿宋_GB2312'
                            run.font.size = Pt(16)
                            run.font.bold = True
            
            # 添加数据行
            for line in data_lines:
                cells_data = [cell.strip() for cell in line.split('|') if cell.strip()]
                if cells_data:
                    row_cells = table.add_row().cells
                    for i, cell_data in enumerate(cells_data):
                        if i < len(row_cells):
                            row_cells[i].text = cell_data
                            # 设置单元格字体
                            for paragraph in row_cells[i].paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = '仿宋_GB2312'
                                    run.font.size = Pt(16)
            
        except Exception as e:
            logger.error(f"创建表格时发生错误: {str(e)}")
    
    def _add_page_numbers(self):
        """
        为整个文档添加连续页码
        """
        try:
            # 获取文档的第一个节
            section = self.document.sections[0]
            
            # 获取页脚
            footer = section.footer
            
            # 清空现有页脚内容
            footer._element.clear()
            
            # 创建页脚段落
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加页码字段
            run = footer_para.add_run()
            
            # 创建页码字段的XML
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            # 添加到run中
            run._element.append(fldChar1)
            run._element.append(instrText)
            run._element.append(fldChar2)
            
            # 设置页码字体
            run.font.name = '仿宋_GB2312'
            run.font.size = Pt(14)
            
            logger.info("页码添加完成")
            
        except Exception as e:
            logger.error(f"添加页码时发生错误: {str(e)}")
    
    def _merge_docx_files(self, main_content: str, attachments: List[Dict]) -> bytes:
        """
        直接拼接docx文件的备选方案
        将正文和附件的docx文件直接拼接成一个新的docx文档
        
        Args:
            main_content: 主要内容
            attachments: 附件列表
        
        Returns:
            bytes: 拼接后的docx文件内容
        """
        try:
            logger.info("开始使用docx直接拼接方案")
            
            # 创建主文档
            merged_doc = Document()
            
            # 添加正文内容
            self._add_main_content_to_doc(merged_doc, main_content)
            
            # 添加附件引用
            self._add_attachment_references_to_doc(merged_doc, attachments)
            
            # 添加每个附件的内容
            for i, attachment in enumerate(attachments, 1):
                if attachment.get('type') == 'word' and 'content' in attachment:
                    # 添加分页符
                    merged_doc.add_page_break()
                    
                    # 添加附件标题
                    self._add_attachment_header_to_doc(merged_doc, i)
                    
                    # 直接拼接Word文档内容
                    self._merge_word_content_to_doc(merged_doc, attachment['content'])
                else:
                    # 其他类型的附件，使用markdown内容
                    merged_doc.add_page_break()
                    self._add_attachment_header_to_doc(merged_doc, i)
                    
                    # 添加附件标题
                    attachment_title = attachment.get('extracted_title', 
                                                    attachment.get('title', 
                                                                 attachment.get('name', f'附件{i}')))
                    self._add_attachment_title_to_doc(merged_doc, attachment_title)
                    
                    # 添加内容
                    content = attachment.get('markdown_content', '')
                    self._add_text_content_to_doc(merged_doc, content)
            
            # 保存到内存
            doc_buffer = BytesIO()
            merged_doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            logger.info("docx直接拼接完成")
            return doc_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"docx直接拼接失败: {str(e)}")
            raise
    
    def _add_main_content_to_doc(self, doc: Document, content: str):
        """向文档添加主要内容"""
        try:
            # 解析内容并添加到文档
            formatted_content = text_processor.format_content_for_document(content)
            
            for item in formatted_content:
                if item['type'] == 'header1':
                    # 一级标题
                    para = doc.add_paragraph(style='Heading1')
                    run = para.add_run(item['text'])
                    run.font.name = '黑体'
                    run.font.size = Pt(22)
                    run.font.bold = True
                    
                elif item['type'] == 'header2':
                    # 二级标题
                    para = doc.add_paragraph(style='Heading2')
                    run = para.add_run(item['text'])
                    run.font.name = '楷体_GB2312'
                    run.font.size = Pt(16)
                    
                elif item['type'] == 'header3':
                    # 三级标题
                    para = doc.add_paragraph(style='Heading3')
                    run = para.add_run(item['text'])
                    run.font.name = '仿宋_GB2312'
                    run.font.size = Pt(16)
                    run.font.bold = True
                    
                else:
                    # 普通段落
                    para = doc.add_paragraph()
                    run = para.add_run(item['text'])
                    run.font.name = '仿宋_GB2312'
                    run.font.size = Pt(16)
                    
        except Exception as e:
            logger.error(f"添加主要内容时发生错误: {str(e)}")
    
    def _add_attachment_references_to_doc(self, doc: Document, attachments: List[Dict]):
        """向文档添加附件引用"""
        if not attachments:
            return
        
        # 添加空行
        doc.add_paragraph()
        
        # 添加附件引用
        if len(attachments) == 1:
            # 单个附件
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Pt(32)
            
            run = para.add_run("附件：附件1、")
            run.font.name = "仿宋_GB2312"
            run.font.size = Pt(16)
            run.font.bold = True
            
            title = attachments[0].get('extracted_title', attachments[0].get('title', '附件'))
            title_run = para.add_run(title)
            title_run.font.name = "仿宋_GB2312"
            title_run.font.size = Pt(16)
        else:
            # 多个附件
            # 第一行：附件：
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Pt(32)
            run = para.add_run("附件：")
            run.font.name = "仿宋_GB2312"
            run.font.size = Pt(16)
            run.font.bold = True
            
            # 每个附件占一行
            for i, attachment in enumerate(attachments, 1):
                para = doc.add_paragraph()
                para.paragraph_format.first_line_indent = Pt(64)
                
                # 附件序号
                number_run = para.add_run(f"附件{i}、")
                number_run.font.name = "仿宋_GB2312"
                number_run.font.size = Pt(16)
                number_run.font.bold = True
                
                # 附件标题
                title = attachment.get('extracted_title', attachment.get('title', f'附件{i}'))
                title_run = para.add_run(title)
                title_run.font.name = "仿宋_GB2312"
                title_run.font.size = Pt(16)
    
    def _add_attachment_header_to_doc(self, doc: Document, attachment_number: int):
        """向文档添加附件标题"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        if attachment_number == 1:
            text = "附件1"
        else:
            text = f"附件{attachment_number}"
            
        run = para.add_run(text)
        run.font.name = "黑体"
        run.font.size = Pt(16)
        run.font.bold = True
    
    def _add_attachment_title_to_doc(self, doc: Document, title: str):
        """向文档添加附件标题"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = para.add_run(title)
        run.font.name = "方正小标宋简体"
        run.font.size = Pt(22)
        run.font.bold = True
    
    def _add_text_content_to_doc(self, doc: Document, content: str):
        """向文档添加文本内容"""
        if not content:
            return
            
        # 简单的文本处理，按行添加
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                para = doc.add_paragraph()
                run = para.add_run(line.strip())
                run.font.name = "仿宋_GB2312"
                run.font.size = Pt(16)
    
    def _merge_word_content_to_doc(self, target_doc: Document, word_content: bytes):
        """将Word文档内容合并到目标文档"""
        try:
            source_doc = Document(BytesIO(word_content))
            
            # 复制所有段落
            for para in source_doc.paragraphs:
                if para.text.strip():
                    new_para = target_doc.add_paragraph()
                    new_para.alignment = para.alignment
                    
                    for run in para.runs:
                        if run.text:
                            new_run = new_para.add_run(run.text)
                            # 复制字体格式
                            if run.font.name:
                                new_run.font.name = run.font.name
                            if run.font.size:
                                new_run.font.size = run.font.size
                            if run.font.bold is not None:
                                new_run.font.bold = run.font.bold
                            if run.font.italic is not None:
                                new_run.font.italic = run.font.italic
            
            # 复制所有表格
            for table in source_doc.tables:
                self._copy_table_to_doc(target_doc, table)
                
        except Exception as e:
            logger.error(f"合并Word内容时发生错误: {str(e)}")
    
    def _copy_table_to_doc(self, target_doc: Document, source_table):
        """将表格复制到目标文档"""
        try:
            # 获取表格尺寸
            rows = len(source_table.rows)
            cols = len(source_table.columns) if source_table.columns else len(source_table.rows[0].cells)
            
            # 创建新表格
            new_table = target_doc.add_table(rows=rows, cols=cols)
            new_table.style = 'Table Grid'
            
            # 复制内容
            for i, source_row in enumerate(source_table.rows):
                for j, source_cell in enumerate(source_row.cells):
                    if j < len(new_table.rows[i].cells):
                        target_cell = new_table.rows[i].cells[j]
                        target_cell.text = source_cell.text
                        
        except Exception as e:
            logger.error(f"复制表格时发生错误: {str(e)}")

# 全局文档生成器实例
document_generator = OfficialDocumentGenerator() 