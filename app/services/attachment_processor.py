"""
附件处理服务 - 处理 Dify 上传的文件
"""
import logging
import requests
import os
import tempfile
from typing import List, Dict, Optional, Tuple
from io import BytesIO
import pandas as pd
from docx import Document
import json
import re
import docx2txt

from app.models.request_models import DifyAttachmentModel

logger = logging.getLogger(__name__)

class AttachmentProcessor:
    """附件处理器"""
    
    def __init__(self):
        """初始化附件处理器"""
        self.supported_extensions = {
            '.docx': 'word',
            '.doc': 'word', 
            '.csv': 'csv',
            '.xlsx': 'excel',
            '.xls': 'excel',
            '.txt': 'text',
            '.md': 'markdown'
        }
    
    def _get_file_type(self, filename: str) -> str:
        """
        根据文件名获取文件类型
        
        Args:
            filename: 文件名
            
        Returns:
            str: 文件类型
        """
        if not filename:
            return 'unknown'
        
        # 获取文件扩展名
        extension = '.' + filename.split('.')[-1].lower() if '.' in filename else ''
        
        return self.supported_extensions.get(extension, 'unknown')
    
    def _extract_title_from_word(self, file_content: bytes, filename: str) -> str:
        """
        从Word文档中智能提取标题
        
        Args:
            file_content: Word文件内容
            filename: 文件名
            
        Returns:
            str: 提取的标题
        """
        try:
            # 检查文件格式
            if filename.lower().endswith('.docx'):
                # DOCX格式，使用python-docx
                file_stream = BytesIO(file_content)
                doc = Document(file_stream)
                
                # 智能提取标题 - 支持多段落标题组合
                title_parts = []
                potential_title_paragraphs = []
                
                # 获取前几个段落，寻找标题
                for i, paragraph in enumerate(doc.paragraphs[:15]):  # 检查前15个段落
                    text = paragraph.text.strip()
                    if text:
                        potential_title_paragraphs.append((i, text, paragraph))
                
                # 尝试智能组合标题
                for i, (para_idx, text, para) in enumerate(potential_title_paragraphs):
                    # 检查是否是标题的开始部分
                    if self._is_title_start(text):
                        title_parts = [text]
                        
                        # 检查后续段落是否是标题的延续
                        for j in range(i + 1, min(i + 5, len(potential_title_paragraphs))):  # 最多检查后续4个段落
                            next_para_idx, next_text, next_para = potential_title_paragraphs[j]
                            
                            # 如果下一个段落是标题的延续部分
                            if self._is_title_continuation(next_text, title_parts):
                                title_parts.append(next_text)
                            else:
                                break
                        
                        # 组合标题
                        if title_parts:
                            combined_title = ''.join(title_parts)
                            if self._is_likely_title(combined_title):
                                logger.info(f"从DOCX文档中提取组合标题: {combined_title}")
                                return combined_title
                
                # 如果组合标题失败，尝试单独的标题
                for para_idx, text, para in potential_title_paragraphs:
                    if self._is_likely_title(text):
                        logger.info(f"从DOCX文档中提取单独标题: {text}")
                        return text
            
            elif filename.lower().endswith('.doc'):
                # DOC格式，使用docx2txt
                with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_file:
                    temp_file.write(file_content)
                    temp_file.flush()
                    
                    try:
                        # 使用docx2txt提取文本
                        text_content = docx2txt.process(temp_file.name)
                        
                        # 按行分割，寻找标题
                        lines = text_content.split('\n')
                        for line in lines[:20]:  # 检查前20行
                            line = line.strip()
                            if line:
                                # 检查是否是标题格式
                                if self._is_likely_title(line):
                                    logger.info(f"从DOC文档中提取标题: {line}")
                                    return line
                    finally:
                        # 清理临时文件
                        try:
                            os.unlink(temp_file.name)
                        except:
                            pass
            
            # 如果没有找到合适的标题，使用文件名（去掉扩展名）
            title_from_filename = filename.rsplit('.', 1)[0] if '.' in filename else filename
            logger.info(f"未找到Word文档标题，使用文件名: {title_from_filename}")
            return title_from_filename
            
        except Exception as e:
            logger.error(f"提取Word文档标题时发生错误: {str(e)}")
            # 出错时返回文件名
            return filename.rsplit('.', 1)[0] if '.' in filename else filename
    
    def _is_title_start(self, text: str) -> bool:
        """
        判断文本是否可能是标题的开始部分
        
        Args:
            text: 文本内容
            
        Returns:
            bool: 是否可能是标题开始
        """
        # 标题开始的常见模式
        title_start_patterns = [
            r'^关于.*',  # 关于...
            r'^.*通知$',  # ...通知
            r'^.*办法$',  # ...办法
            r'^.*方案$',  # ...方案
            r'^.*规定$',  # ...规定
            r'^.*制度$',  # ...制度
            r'^.*报告$',  # ...报告
            r'^.*情况$',  # ...情况
            r'^.*统计.*',  # ...统计...
            r'^.*名单.*',  # ...名单...
            r'^.*清单.*',  # ...清单...
            r'^第.*届.*',  # 第X届...
            r'^.*30强.*',  # ...30强...
            r'^.*全国文化企业.*',  # ...全国文化企业...
        ]
        
        return any(re.search(pattern, text) for pattern in title_start_patterns)
    
    def _is_title_continuation(self, text: str, title_parts: List[str]) -> bool:
        """
        判断文本是否是标题的延续部分
        
        Args:
            text: 当前文本
            title_parts: 已有的标题部分
            
        Returns:
            bool: 是否是标题延续
        """
        if not title_parts:
            return False
        
        # 获取前一个标题部分
        previous_part = title_parts[-1]
        
        # 标题延续的特征
        continuation_patterns = [
            r'^有关.*的.*',  # 有关...的...
            r'^.*的通知$',  # ...的通知
            r'^.*的办法$',  # ...的办法
            r'^.*的方案$',  # ...的方案
            r'^.*的规定$',  # ...的规定
            r'^.*事项.*',  # ...事项...
            r'^.*工作.*',  # ...工作...
            r'^.*名单.*',  # ...名单...
            r'^.*情况.*',  # ...情况...
            r'^.*及分布情况$',  # ...及分布情况
            r'^.*30强.*',  # ...30强...
            r'^.*全国成长性.*',  # ...全国成长性...
        ]
        
        # 检查是否匹配延续模式
        matches_continuation = any(re.search(pattern, text) for pattern in continuation_patterns)
        
        # 检查长度是否合理（延续部分通常不会太长）
        reasonable_length = len(text) <= 50
        
        # 检查是否是明显的正文开始
        content_start_indicators = [
            '根据', '按照', '为了', '现将', '现印发', '请', '各单位',
            '各部门', '认真', '贯彻', '执行', '落实'
        ]
        is_content_start = any(text.startswith(indicator) for indicator in content_start_indicators)
        
        return matches_continuation and reasonable_length and not is_content_start

    def _is_likely_title(self, text: str) -> bool:
        """
        判断文本是否可能是标题
        
        Args:
            text: 文本内容
            
        Returns:
            bool: 是否可能是标题
        """
        # 标题通常的特征：
        # 1. 长度适中（不会太短也不会太长）
        # 2. 不以标点符号结尾（除了特殊情况）
        # 3. 包含关键词
        # 4. 格式特征
        
        # 长度检查 - 放宽长度限制以适应更长的标题
        if len(text) < 5 or len(text) > 150:
            return False
        
        # 检查是否包含标题关键词 - 根据用户要求扩展关键词
        title_keywords = [
            # 统计表、清单、名单、汇总表类
            '统计表', '清单', '名单', '汇总表', '汇总清单', '统计', '汇总', '分布情况',
            '考核统计', '代表名单', '参会代表', '绩效考核', '年度统计',
            
            # 上级来文类
            '通知', '函', '意见', '决定', '批复', '指示', '要求', '部署',
            '关于', '省教育厅', '市政府', '区政府', '教办', '政办',
            
            # 制度办法方案类
            '办法', '规定', '制度', '方案', '细则', '标准', '规范', '程序',
            '管理办法', '实施办法', '工作方案', '实施方案', '试行',
            
            # 合同协议类
            '合同', '协议', '服务合同', '技术服务', '采购合同', '编号',
            
            # 其他常见标题词
            '报告', '情况', '工作', '实施', '管理', '企业', '公司', '评选', '评审',
            '总结', '计划', '安排', '部署', '要点', '措施', '建议', '意见',
            '全国文化企业', '30强', '成长性', '分布', '届'
        ]
        
        has_title_keyword = any(keyword in text for keyword in title_keywords)
        
        # 检查格式特征 - 扩展标题格式识别
        title_patterns = [
            r'第[一二三四五六七八九十\d]+届',  # 第X届
            r'关于.*的.*',  # 关于...的...
            r'.*情况.*',  # ...情况...
            r'.*报告.*',  # ...报告...
            r'.*通知.*',  # ...通知...
            r'.*办法.*',  # ...办法...
            r'.*方案.*',  # ...方案...
            r'.*制度.*',  # ...制度...
            r'.*规定.*',  # ...规定...
            r'.*名单.*',  # ...名单...
            r'.*统计.*',  # ...统计...
            r'.*清单.*',  # ...清单...
            r'.*合同.*',  # ...合同...
            r'.*协议.*',  # ...协议...
            r'.*〔\d+〕\d+号',  # 公文编号格式
            r'.*（试行）',  # 试行文件
            r'.*（编号.*）',  # 编号格式
            r'.*".*".*',  # 引号包含的内容
            r'.*年度.*',  # 年度相关
            r'.*工作.*',  # 工作相关
            r'.*企业.*强.*',  # 企业强相关
        ]
        
        has_title_pattern = any(re.search(pattern, text) for pattern in title_patterns)
        
        # 检查是否以常见的非标题结尾
        bad_endings = ['。', '！', '？', '：', '；', '，', '、']
        ends_badly = any(text.endswith(ending) for ending in bad_endings)
        
        # 检查是否是明显的正文内容
        content_indicators = [
            '根据', '按照', '为了', '现将', '现印发', '请', '各单位',
            '各部门', '认真', '贯彻', '执行', '落实', '具体如下',
            '现就', '经研究', '决定', '同意', '批准'
        ]
        
        is_content = any(text.startswith(indicator) for indicator in content_indicators)
        
        # 综合判断
        score = 0
        if has_title_keyword:
            score += 2
        if has_title_pattern:
            score += 2
        if not ends_badly:
            score += 1
        if not is_content:
            score += 1
        
        # 特殊情况：如果包含"全国文化企业30强"等特定关键词，直接认为是标题
        special_keywords = [
            '全国文化企业', '30强', '成长性', '分布情况', '名单', 
            '第十六届', '全国成长性文化企业', '名单及分布情况'
        ]
        if any(keyword in text for keyword in special_keywords):
            score += 3
        
        # 如果文本很长且包含多个关键词，也可能是完整标题
        if len(text) > 30 and has_title_keyword and has_title_pattern:
            score += 2
        
        return score >= 3

    def process_dify_attachments(self, dify_attachments: List[DifyAttachmentModel]) -> List[Dict]:
        """
        处理 Dify 上传的附件
        
        Args:
            dify_attachments: Dify 附件列表
            
        Returns:
            List[Dict]: 处理后的附件列表
        """
        processed_attachments = []
        
        for attachment in dify_attachments:
            try:
                logger.info(f"开始处理附件: {attachment.name}")
                
                # 下载文件内容
                file_content = self._download_file_from_url(attachment.url)
                if not file_content:
                    logger.error(f"无法下载附件: {attachment.name}")
                    continue
                
                # 获取文件类型
                file_type = self._get_file_type(attachment.name)
                logger.info(f"文件类型: {file_type}")
                
                # 处理不同类型的文件
                processed_attachment = {
                    'name': attachment.name,
                    'type': file_type,
                    'original_url': attachment.url,
                    'content': file_content  # 保存原始内容用于Word拼接
                }
                
                if file_type == 'word':
                    # 提取智能标题
                    smart_title = self._extract_title_from_word(file_content, attachment.name)
                    processed_attachment['title'] = smart_title
                    processed_attachment['extracted_title'] = smart_title  # 添加extracted_title字段
                    processed_attachment['markdown_content'] = self._process_word_file(file_content)
                elif file_type == 'csv':
                    # 为非Word文件设置extracted_title字段
                    file_title = attachment.name.rsplit('.', 1)[0] if '.' in attachment.name else attachment.name
                    processed_attachment['title'] = file_title
                    processed_attachment['extracted_title'] = file_title
                    processed_attachment['markdown_content'] = self._process_csv_file(file_content)
                elif file_type == 'excel':
                    file_title = attachment.name.rsplit('.', 1)[0] if '.' in attachment.name else attachment.name
                    processed_attachment['title'] = file_title
                    processed_attachment['extracted_title'] = file_title
                    processed_attachment['markdown_content'] = self._process_excel_file(file_content)
                elif file_type == 'text':
                    file_title = attachment.name.rsplit('.', 1)[0] if '.' in attachment.name else attachment.name
                    processed_attachment['title'] = file_title
                    processed_attachment['extracted_title'] = file_title
                    processed_attachment['markdown_content'] = self._process_text_file(file_content)
                elif file_type == 'markdown':
                    file_title = attachment.name.rsplit('.', 1)[0] if '.' in attachment.name else attachment.name
                    processed_attachment['title'] = file_title
                    processed_attachment['extracted_title'] = file_title
                    processed_attachment['markdown_content'] = self._process_markdown_file(file_content)
                else:
                    logger.warning(f"不支持的文件类型: {file_type}")
                    file_title = attachment.name.rsplit('.', 1)[0] if '.' in attachment.name else attachment.name
                    processed_attachment['title'] = file_title
                    processed_attachment['extracted_title'] = file_title
                    processed_attachment['markdown_content'] = f"不支持的文件类型: {file_type}"
                
                processed_attachments.append(processed_attachment)
                logger.info(f"附件处理完成: {attachment.name}")
                
            except Exception as e:
                logger.error(f"处理附件 {attachment.name} 时发生错误: {str(e)}")
                # 添加错误信息的附件
                error_title = attachment.name.rsplit('.', 1)[0] if '.' in attachment.name else attachment.name
                processed_attachments.append({
                    'name': attachment.name,
                    'type': 'error',
                    'title': error_title,
                    'extracted_title': error_title,
                    'markdown_content': f"处理附件时发生错误: {str(e)}"
                })
        
        return processed_attachments

    def _download_file_from_url(self, url: str) -> Optional[bytes]:
        """
        从URL下载文件内容
        
        Args:
            url: 文件URL
            
        Returns:
            Optional[bytes]: 文件内容，如果失败则返回None
        """
        try:
            # 处理 data URL
            if url.startswith('data:'):
                return self._process_data_url(url)
            
            # 处理普通 HTTP/HTTPS URL
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            return response.content
            
        except Exception as e:
            logger.error(f"下载文件时发生错误: {str(e)}")
            return None
    
    def _process_data_url(self, data_url: str) -> Optional[bytes]:
        """
        处理 data URL 格式的文件
        
        Args:
            data_url: data URL
            
        Returns:
            Optional[bytes]: 解码后的文件内容
        """
        try:
            import base64
            
            # data URL 格式: data:[<mediatype>][;base64],<data>
            if ',' not in data_url:
                logger.error("无效的 data URL 格式")
                return None
            
            header, data = data_url.split(',', 1)
            
            # 检查是否是 base64 编码
            if 'base64' in header:
                return base64.b64decode(data)
            else:
                # 如果不是 base64，假设是 URL 编码
                from urllib.parse import unquote
                return unquote(data).encode('utf-8')
                
        except Exception as e:
            logger.error(f"处理 data URL 时发生错误: {str(e)}")
            return None
    
    def _process_word_file(self, file_content: bytes) -> str:
        """
        处理 Word 文件，提取文本和表格
        
        Args:
            file_content: Word 文件内容
            
        Returns:
            str: 提取的 markdown 内容
        """
        try:
            # 使用 BytesIO 创建文件对象
            file_stream = BytesIO(file_content)
            doc = Document(file_stream)
            
            markdown_content = []
            
            # 提取段落
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    markdown_content.append(text)
            
            # 提取表格
            for table in doc.tables:
                table_markdown = self._convert_table_to_markdown(table)
                if table_markdown:
                    markdown_content.append(table_markdown)
            
            return '\n\n'.join(markdown_content)
            
        except Exception as e:
            logger.error(f"处理 Word 文件时发生错误: {str(e)}")
            return f"Word 文件处理失败: {str(e)}"
    
    def _process_csv_file(self, file_content: bytes) -> str:
        """
        处理 CSV 文件
        
        Args:
            file_content: CSV 文件内容
            
        Returns:
            str: 转换后的 markdown 表格
        """
        try:
            # 尝试不同的编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-8-sig']
            
            for encoding in encodings:
                try:
                    csv_text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                # 如果所有编码都失败，使用 utf-8 并忽略错误
                csv_text = file_content.decode('utf-8', errors='ignore')
            
            # 使用 pandas 读取 CSV
            from io import StringIO
            df = pd.read_csv(StringIO(csv_text))
            
            # 转换为 markdown 表格
            return self._dataframe_to_markdown(df)
            
        except Exception as e:
            logger.error(f"处理 CSV 文件时发生错误: {str(e)}")
            return f"CSV 文件处理失败: {str(e)}"
    
    def _process_excel_file(self, file_content: bytes) -> str:
        """
        处理 Excel 文件
        
        Args:
            file_content: Excel 文件内容
            
        Returns:
            str: 转换后的 markdown 表格
        """
        try:
            file_stream = BytesIO(file_content)
            df = pd.read_excel(file_stream)
            
            return self._dataframe_to_markdown(df)
            
        except Exception as e:
            logger.error(f"处理 Excel 文件时发生错误: {str(e)}")
            return f"Excel 文件处理失败: {str(e)}"
    
    def _process_text_file(self, file_content: bytes) -> str:
        """
        处理纯文本文件
        
        Args:
            file_content: 文本文件内容
            
        Returns:
            str: 文本内容
        """
        try:
            # 尝试不同的编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-8-sig']
            
            for encoding in encodings:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # 如果所有编码都失败，使用 utf-8 并忽略错误
            return file_content.decode('utf-8', errors='ignore')
            
        except Exception as e:
            logger.error(f"处理文本文件时发生错误: {str(e)}")
            return f"文本文件处理失败: {str(e)}"
    
    def _process_markdown_file(self, file_content: bytes) -> str:
        """
        处理 Markdown 文件
        
        Args:
            file_content: Markdown 文件内容
            
        Returns:
            str: Markdown 内容
        """
        return self._process_text_file(file_content)
    
    def _convert_table_to_markdown(self, table) -> str:
        """
        将 Word 表格转换为 Markdown 表格
        
        Args:
            table: Word 表格对象
            
        Returns:
            str: Markdown 表格
        """
        try:
            rows = []
            
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    cells.append(cell_text)
                rows.append(cells)
            
            if not rows:
                return ""
            
            # 构建 Markdown 表格
            markdown_lines = []
            
            # 表头
            header = "| " + " | ".join(rows[0]) + " |"
            markdown_lines.append(header)
            
            # 分隔线
            separator = "| " + " | ".join(["---"] * len(rows[0])) + " |"
            markdown_lines.append(separator)
            
            # 数据行
            for row in rows[1:]:
                if len(row) == len(rows[0]):  # 确保列数一致
                    data_row = "| " + " | ".join(row) + " |"
                    markdown_lines.append(data_row)
            
            return "\n".join(markdown_lines)
            
        except Exception as e:
            logger.error(f"转换表格时发生错误: {str(e)}")
            return ""
    
    def _dataframe_to_markdown(self, df: pd.DataFrame) -> str:
        """
        将 DataFrame 转换为 Markdown 表格
        
        Args:
            df: pandas DataFrame
            
        Returns:
            str: Markdown 表格
        """
        try:
            # 处理空值
            df = df.fillna('')
            
            # 转换为字符串
            df = df.astype(str)
            
            # 构建 Markdown 表格
            markdown_lines = []
            
            # 表头
            headers = list(df.columns)
            header = "| " + " | ".join(headers) + " |"
            markdown_lines.append(header)
            
            # 分隔线
            separator = "| " + " | ".join(["---"] * len(headers)) + " |"
            markdown_lines.append(separator)
            
            # 数据行
            for _, row in df.iterrows():
                data_row = "| " + " | ".join(row.values) + " |"
                markdown_lines.append(data_row)
            
            return "\n".join(markdown_lines)
            
        except Exception as e:
            logger.error(f"转换 DataFrame 时发生错误: {str(e)}")
            return f"表格转换失败: {str(e)}"
    
    def _get_attachment_type(self, extension: str) -> str:
        """
        根据文件扩展名获取附件类型
        
        Args:
            extension: 文件扩展名
            
        Returns:
            str: 附件类型
        """
        type_mapping = {
            '.docx': 'word',
            '.doc': 'word',
            '.csv': 'csv',
            '.xlsx': 'excel',
            '.xls': 'excel',
            '.txt': 'text',
            '.md': 'markdown'
        }
        
        return type_mapping.get(extension.lower(), 'unknown')
    
    def _clean_filename(self, filename: str) -> str:
        """
        清理文件名，移除扩展名和特殊字符
        
        Args:
            filename: 原始文件名
            
        Returns:
            str: 清理后的文件名
        """
        # 移除扩展名
        name_without_ext = filename.rsplit('.', 1)[0]
        
        # 移除时间戳前缀（如果存在）
        # 例如：20250711_085907_测试报告.docx -> 测试报告
        if re.match(r'^\d{8}_\d{6}_', name_without_ext):
            name_without_ext = name_without_ext[16:]  # 移除前16个字符
        
        return name_without_ext.strip()

# 全局附件处理器实例
attachment_processor = AttachmentProcessor() 